from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "software-documentation.md"
OUTPUT = ROOT / "docs" / "Studio-Ledger-Software-Documentation.docx"
HERO = ROOT / "public" / "og.png"

NAVY = "1F3A5F"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5E6875"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "C7D0DA"
WHITE = "FFFFFF"
INK = "20262E"
GOLD = "9A6A16"


def set_run_font(run, name="Calibri", size=11, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "left" if edge == "start" else "right" if edge == "end" else edge
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def add_numbering_definition(doc, ordered):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if ordered else "•")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(fonts)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])
    p_pr.append(num_pr)


def add_inline(paragraph, text, size=11, color=INK, bold=False, italic=False):
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    for piece in pattern.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            run = paragraph.add_run(piece[2:-2])
            set_run_font(run, size=size, color=color, bold=True, italic=italic)
        elif piece.startswith("`") and piece.endswith("`"):
            run = paragraph.add_run(piece[1:-1])
            set_run_font(run, name="Courier New", size=max(8.5, size - 0.5), color=DARK_BLUE)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), LIGHT_GRAY)
            run._r.get_or_add_rPr().append(shading)
        else:
            run = paragraph.add_run(piece)
            set_run_font(run, size=size, color=color, bold=bold, italic=italic)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for name in ("Caption",):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(9)
        style.font.color.rgb = RGBColor.from_string(MUTED)
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(4)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    add_inline(hp, "STUDIO LEDGER  /  SOFTWARE DOCUMENTATION", size=8.5, color=MUTED, bold=True)
    p_pr = hp._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), BORDER)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    run = fp.add_run("Studio Ledger  •  Version 2.0  •  ")
    set_run_font(run, size=9, color=MUTED)
    add_page_field(fp)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("PRODUCT & OPERATIONS MANUAL")
    set_run_font(r, size=10, color=GOLD, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Studio Ledger")
    set_run_font(r, size=30, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("Software Documentation")
    set_run_font(r, size=17, color=DARK_BLUE)

    if HERO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(16)
        picture = p.add_run().add_picture(str(HERO), width=Inches(6.25))
        doc_pr = picture._inline.docPr
        doc_pr.set("descr", "Interior design plan, material samples, ruler, and project ledger")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    add_inline(p, "As-built functional and technical reference", size=12, color=MUTED, italic=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    add_inline(p, "Version 2.0  •  28 July 2026", size=10.5, color=MUTED, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    add_inline(p, "For studio owners, operations, finance, project teams, developers, and administrators", size=9.5, color=MUTED)
    doc.add_page_break()


def add_contents(doc, headings):
    p = doc.add_paragraph("Contents", style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    for heading in headings:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        p.paragraph_format.space_after = Pt(3)
        add_inline(p, heading, size=10.5, color=DARK_BLUE)
    doc.add_page_break()


def table_widths(headers):
    count = len(headers)
    if count == 2:
        return [2700, 6660]
    if count == 3:
        return [2200, 3380, 3780]
    if count == 4:
        if headers and headers[0].lower() == "area":
            return [1450, 2650, 2250, 3010]
        return [1650, 2570, 2570, 2570]
    return [9360 // count] * count


def add_table(doc, rows):
    headers = rows[0]
    widths = table_widths(headers)
    table = doc.add_table(rows=len(rows), cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for table_row in table.rows:
        set_row_cant_split(table_row)
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.rows[row_index].cells[col_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                set_cell_fill(cell, LIGHT_BLUE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline(p, value, size=9.1 if len(headers) >= 4 else 9.5,
                       color=NAVY if row_index == 0 else INK,
                       bold=row_index == 0)
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(2)


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.05
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GRAY)
    p_pr.append(shd)
    run = p.add_run("\n".join(lines))
    set_run_font(run, name="Courier New", size=8.7, color=DARK_BLUE)


def parse_markdown(doc, markdown, bullet_num_id):
    lines = markdown.splitlines()
    first_rule_seen = False
    active_decimal_num_id = None
    index = 0
    while index < len(lines):
        raw = lines[index]
        line = raw.rstrip()
        if not line:
            index += 1
            continue
        if line.startswith("# "):
            index += 1
            continue
        if line.startswith("## "):
            heading = line[3:].strip()
            if heading != "Software Documentation":
                paragraph = doc.add_paragraph(heading, style="Heading 1")
                if heading == "4. User roles and responsibilities":
                    paragraph.paragraph_format.page_break_before = True
            active_decimal_num_id = None
            index += 1
            continue
        if line.startswith("### "):
            doc.add_paragraph(line[4:].strip(), style="Heading 2")
            active_decimal_num_id = None
            index += 1
            continue
        if line.startswith("#### "):
            doc.add_paragraph(line[5:].strip(), style="Heading 3")
            active_decimal_num_id = None
            index += 1
            continue
        if line == "---":
            if not first_rule_seen:
                first_rule_seen = True
            index += 1
            continue
        if line.startswith("```"):
            code = []
            index += 1
            while index < len(lines) and not lines[index].startswith("```"):
                code.append(lines[index])
                index += 1
            index += 1
            add_code_block(doc, code)
            active_decimal_num_id = None
            continue
        if line.startswith("|") and index + 1 < len(lines) and re.match(r"^\|?[\s:|-]+\|", lines[index + 1]):
            rows = []
            header = [cell.strip() for cell in line.strip("|").split("|")]
            rows.append(header)
            index += 2
            while index < len(lines) and lines[index].startswith("|"):
                rows.append([cell.strip() for cell in lines[index].strip("|").split("|")])
                index += 1
            add_table(doc, rows)
            active_decimal_num_id = None
            continue
        bullet_match = re.match(r"^\s*-\s+(.+)$", line)
        number_match = re.match(r"^\s*\d+\.\s+(.+)$", line)
        if bullet_match or number_match:
            text = (bullet_match or number_match).group(1)
            index += 1
            while index < len(lines):
                continuation = lines[index]
                indentation = len(continuation) - len(continuation.lstrip())
                is_new_list_item = re.match(r"^\s*(?:-|\d+\.)\s+", continuation)
                if continuation.strip() and indentation > 0 and not is_new_list_item:
                    text += " " + continuation.strip()
                    index += 1
                    continue
                break
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.25
            if number_match and active_decimal_num_id is None:
                active_decimal_num_id = add_numbering_definition(doc, ordered=True)
            apply_numbering(p, bullet_num_id if bullet_match else active_decimal_num_id)
            add_inline(p, text)
            continue
        active_decimal_num_id = None
        if re.match(r"^\*\*[^*]+:\*\*", line):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            add_inline(p, line)
            index += 1
            continue

        parts = [line.strip()]
        index += 1
        while index < len(lines):
            nxt = lines[index].rstrip()
            if not nxt or nxt.startswith(("#", "|", "```", "- ")) or re.match(r"^\s*\d+\.\s+", nxt) or nxt == "---":
                break
            parts.append(nxt.strip())
            index += 1
        paragraph_text = " ".join(parts)
        p = doc.add_paragraph()
        p.paragraph_format.widow_control = True
        add_inline(p, paragraph_text)


def set_document_properties(doc):
    props = doc.core_properties
    props.title = "Studio Ledger Software Documentation"
    props.subject = "As-built functional and technical product documentation"
    props.author = "Studio Ledger Product Team"
    props.keywords = "Studio Ledger, interior design, procurement, estimates, receivables, payables"
    props.comments = "Generated from the repository documentation source."


def main():
    markdown = SOURCE.read_text(encoding="utf-8")
    headings = [
        line[3:].strip() for line in markdown.splitlines()
        if line.startswith("## ") and line[3:].strip() != "Software Documentation"
    ]
    doc = Document()
    style_document(doc)
    set_document_properties(doc)
    bullet_num_id = add_numbering_definition(doc, ordered=False)
    add_cover(doc)
    add_contents(doc, headings)
    parse_markdown(doc, markdown, bullet_num_id)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    sys.exit(main())
